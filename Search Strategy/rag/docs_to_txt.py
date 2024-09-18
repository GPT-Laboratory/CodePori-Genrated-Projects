from docx import Document

def extract_text_and_tables(docx_path):
    document = Document(docx_path)
    text = []
    tables = []

    for para in document.paragraphs:
        text.append(para.text)

    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return '\n'.join(text), tables

def save_to_txt(text, tables, output_path):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('Extracted Text:\n')
        f.write(text)
        f.write('\n\nExtracted Tables:\n')
        for table in tables:
            for row in table:
                f.write('\t'.join(row) + '\n')
            f.write('\n')
#Path to the input File
docx_path = 'C:/Users/ZRasheed/Desktop/Pekka Task/Marko Task/rag/FEMMa Home-v6-20240520_111308.docx'

# Path to the output text file
output_path = 'context.txt'

# Extract text and tables
text, tables = extract_text_and_tables(docx_path)

# Save the extracted content to a text file
save_to_txt(text, tables, output_path)